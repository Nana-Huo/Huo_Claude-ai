import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
import sys
import io

# 设置输出编码
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def create_resume():
    # 创建新文档
    doc = docx.Document()

    # 设置页面边距
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 添加标题 - 芦若禾
    title = doc.add_paragraph()
    title_run = title.add_run('芦若禾')
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(44, 62, 80)  # 深蓝色
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加求职意向
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('C++开发工程师 | 嵌入式系统开发')
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(52, 73, 94)  # 灰蓝色
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    doc.add_paragraph()

    # 添加个人信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = "电话：15209179155 | 邮箱：3296298836@qq.com | 现居：陕西省宝鸡市"
    info_run = info_para.add_run(info_text)
    info_run.font.name = '微软雅黑'
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(127, 140, 141)  # 灰色

    # 添加分割线
    doc.add_paragraph()

    # 教育背景
    edu_title = doc.add_paragraph()
    edu_title_run = edu_title.add_run('教育背景')
    edu_title_run.font.name = '微软雅黑'
    edu_title_run.font.size = Pt(14)
    edu_title_run.font.bold = True
    edu_title_run.font.color.rgb = RGBColor(231, 76, 60)  # 红色

    # 教育经历
    edu_para = doc.add_paragraph()
    edu_para.add_run('陕西理工大学').font.size = Pt(11)
    edu_para.add_run(' | ').font.size = Pt(11)
    edu_para.add_run('信息与计算科学').font.size = Pt(11)
    edu_para.add_run(' | ').font.size = Pt(11)
    edu_para.add_run('本科').font.size = Pt(11)
    edu_para.add_run(' | ').font.size = Pt(11)
    edu_para.add_run('2022.09 - 2026.06').font.size = Pt(11)

    # 核心课程
    courses_para = doc.add_paragraph()
    courses_para.add_run('核心课程：').font.bold = True
    courses_para.add_run('计算机网络、计算机组成原理、数据结构、数值计算方法、C/C++语言、Java语言、算法与数据分析、数据库应用技术、数学分析、高等代数、解析几何、概率统计')
    courses_para.runs[0].font.size = Pt(9)
    for run in courses_para.runs[1:]:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(75, 75, 75)

    # 成绩信息
    score_para = doc.add_paragraph()
    score_para.add_run('成绩排名：').font.bold = True
    score_para.add_run('专业前10% | 平均绩点：3/4.0')
    score_para.runs[0].font.size = Pt(9)
    score_para.runs[1].font.size = Pt(9)
    score_para.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    # 添加空行
    doc.add_paragraph()

    # 专业技能
    skills_title = doc.add_paragraph()
    skills_title_run = skills_title.add_run('专业技能')
    skills_title_run.font.name = '微软雅黑'
    skills_title_run.font.size = Pt(14)
    skills_title_run.font.bold = True
    skills_title_run.font.color.rgb = RGBColor(231, 76, 60)

    # 技能列表
    skills_list = [
        '熟练掌握C/C++语言，具备扎实的编程基础和算法能力',
        '精通ARM架构嵌入式系统开发，能完成智能语音AI设备的硬件驱动开发、固件调试',
        '熟悉VMware Workstation嵌入式操作系统移植与优化',
        '掌握ASR/TTS算法嵌入式落地，可在STM32/RK3399部署轻量语音模型',
        '具备语音交互软硬件联调经验',
        '熟悉MySQL数据库设计和应用',
        '掌握OpenCV图像处理和人脸识别技术',
        '了解多线程编程和设计模式'
    ]

    for skill in skills_list:
        skill_para = doc.add_paragraph()
        skill_para.add_run('• ').font.size = Pt(9)
        skill_para.add_run(skill).font.size = Pt(9)
        skill_para.paragraph_format.left_indent = Inches(0.25)

    # 添加空行
    doc.add_paragraph()

    # 实习经历
    exp_title = doc.add_paragraph()
    exp_title_run = exp_title.add_run('实习经历')
    exp_title_run.font.name = '微软雅黑'
    exp_title_run.font.size = Pt(14)
    exp_title_run.font.bold = True
    exp_title_run.font.color.rgb = RGBColor(231, 76, 60)

    # 四川讯方信息技术有限公司
    company1 = doc.add_paragraph()
    company1.add_run('四川讯方信息技术有限公司').font.bold = True
    company1.add_run(' | C++开发实习生 | 2025.09')
    company1.runs[0].font.size = Pt(10)
    company1.runs[1].font.size = Pt(9)
    company1.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    work1_para = doc.add_paragraph()
    work1_para.add_run('• 开发C++仓储管理系统，集成MySQL数据库存储和管理，实现商品出入库等核心功能\n')
    work1_para.add_run('• 集成OpenCV人脸识别技术，提升系统安全性和用户体验\n')
    work1_para.add_run('• 运用多线程技术优化系统性能，确保高并发场景下的稳定性\n')
    work1_para.add_run('• 应用设计模式构建可扩展的系统架构，支持双角色权限管理\n')
    work1_para.add_run('• 以面向对象思想完成系统设计，为智能仓储提供可靠解决方案')
    for run in work1_para.runs:
        run.font.size = Pt(9)
    work1_para.paragraph_format.left_indent = Inches(0.25)

    # 宝鸡市新华书店
    company2 = doc.add_paragraph()
    company2.add_run('宝鸡市新华书店有限责任公司').font.bold = True
    company2.add_run(' | 业务员 | 2024.08')
    company2.runs[0].font.size = Pt(10)
    company2.runs[1].font.size = Pt(9)
    company2.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    work2_para = doc.add_paragraph()
    work2_para.add_run('• 负责日常业务接待和客户服务工作，提升沟通协调能力\n')
    work2_para.add_run('• 协助完成图书整理和库存管理工作，培养细致认真的工作态度')
    for run in work2_para.runs:
        run.font.size = Pt(9)
    work2_para.paragraph_format.left_indent = Inches(0.25)

    # 社区实践
    community_para = doc.add_paragraph()
    community_para.add_run('宝鸡市渭滨区机场街社区').font.bold = True
    community_para.add_run(' | 实践生 | 2023.01')
    community_para.runs[0].font.size = Pt(10)
    community_para.runs[1].font.size = Pt(9)
    community_para.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    # 添加空行
    doc.add_paragraph()

    # 项目经验
    project_title = doc.add_paragraph()
    project_title_run = project_title.add_run('项目经验')
    project_title_run.font.name = '微软雅黑'
    project_title_run.font.size = Pt(14)
    project_title_run.font.bold = True
    project_title_run.font.color.rgb = RGBColor(231, 76, 60)

    # C++仓储管理系统
    project1 = doc.add_paragraph()
    project1.add_run('C++仓储管理系统').font.bold = True
    project1.add_run(' | 核心开发')
    project1.runs[0].font.size = Pt(10)
    project1.runs[1].font.size = Pt(9)
    project1.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    project1_desc = doc.add_paragraph()
    project1_desc.add_run('• 基于C++开发的智能仓储管理解决方案，集成MySQL数据库实现数据持久化\n')
    project1_desc.add_run('• 集成OpenCV人脸识别技术，实现系统安全认证和用户身份验证\n')
    project1_desc.add_run('• 采用多线程编程技术，提升系统并发处理能力和响应速度\n')
    project1_desc.add_run('• 应用工厂模式、单例模式等设计模式，确保系统架构的可扩展性和维护性\n')
    project1_desc.add_run('• 实现管理员和普通用户双角色权限管理，确保系统安全性')
    for run in project1_desc.runs:
        run.font.size = Pt(9)
    project1_desc.paragraph_format.left_indent = Inches(0.25)

    # 添加空行
    doc.add_paragraph()

    # 荣誉奖项
    awards_title = doc.add_paragraph()
    awards_title_run = awards_title.add_run('荣誉奖项')
    awards_title_run.font.name = '微软雅黑'
    awards_title_run.font.size = Pt(14)
    awards_title_run.font.bold = True
    awards_title_run.font.color.rgb = RGBColor(231, 76, 60)

    awards_list = [
        '全国大学生数学建模大赛省级二等奖',
        '中国国际大学生创新大赛校级三等奖',
        '陕西理工大学市场调研与分析大赛校级二等奖',
        '计算机设计大赛校级三等奖',
        '陕西理工大学"一封家书"征文活动二等奖',
        '陕西理工大学红色与古风歌曲传唱大赛二等奖',
        '2024-2025笃行奖学金'
    ]

    for award in awards_list:
        award_para = doc.add_paragraph()
        award_para.add_run('• ').font.size = Pt(9)
        award_para.add_run(award).font.size = Pt(9)
        award_para.paragraph_format.left_indent = Inches(0.25)

    # 添加空行
    doc.add_paragraph()

    # 校内实践
    campus_title = doc.add_paragraph()
    campus_title_run = campus_title.add_run('校内实践')
    campus_title_run.font.name = '微软雅黑'
    campus_title_run.font.size = Pt(14)
    campus_title_run.font.bold = True
    campus_title_run.font.color.rgb = RGBColor(231, 76, 60)

    # 班级生活委员
    campus1 = doc.add_paragraph()
    campus1.add_run('陕西理工大学').font.bold = True
    campus1.add_run(' | 生活委员 | 2022.09 - 2024.09')
    campus1.runs[0].font.size = Pt(10)
    campus1.runs[1].font.size = Pt(9)
    campus1.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    campus1_desc = doc.add_paragraph()
    campus1_desc.add_run('• 协助辅导员处理班级日常事务，负责班级经费管理和活动组织\n')
    campus1_desc.add_run('• 组织班级团建活动，增强班级凝聚力和同学间的沟通交流')
    for run in campus1_desc.runs:
        run.font.size = Pt(9)
    campus1_desc.paragraph_format.left_indent = Inches(0.25)

    # 学生会办公室委员
    campus2 = doc.add_paragraph()
    campus2.add_run('陕西理工大学学生会').font.bold = True
    campus2.add_run(' | 办公室委员 | 2022.09 - 2023.09')
    campus2.runs[0].font.size = Pt(10)
    campus2.runs[1].font.size = Pt(9)
    campus2.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    campus2_desc = doc.add_paragraph()
    campus2_desc.add_run('• 负责学生会日常文档处理和表格制作工作\n')
    campus2_desc.add_run('• 协助组织校园文化活动，提升活动策划和执行能力')
    for run in campus2_desc.runs:
        run.font.size = Pt(9)
    campus2_desc.paragraph_format.left_indent = Inches(0.25)

    # 返家乡社会实践
    social_para = doc.add_paragraph()
    social_para.add_run('"返家乡"社会实践活动').font.bold = True
    social_para.add_run(' | 参与者 | 2022.01')
    social_para.runs[0].font.size = Pt(10)
    social_para.runs[1].font.size = Pt(9)
    social_para.runs[1].font.color.rgb = RGBColor(75, 75, 75)

    # 添加空行
    doc.add_paragraph()

    # 证书技能
    cert_title = doc.add_paragraph()
    cert_title_run = cert_title.add_run('证书技能')
    cert_title_run.font.name = '微软雅黑'
    cert_title_run.font.size = Pt(14)
    cert_title_run.font.bold = True
    cert_title_run.font.color.rgb = RGBColor(231, 76, 60)

    cert_list = [
        '英语四级证书',
        '普通话二级甲等证书',
        'C1驾驶证证书'
    ]

    for cert in cert_list:
        cert_para = doc.add_paragraph()
        cert_para.add_run('• ').font.size = Pt(9)
        cert_para.add_run(cert).font.size = Pt(9)
        cert_para.paragraph_format.left_indent = Inches(0.25)

    # 保存文档
    output_path = r"C:\Users\霍冠华\Desktop\简历\简历修改测试.docx"
    doc.save(output_path)
    print(f"简历已成功生成并保存到: {output_path}")

if __name__ == "__main__":
    create_resume()