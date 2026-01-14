import os
import sys

# 添加路径
sys.path.append(r'C:\Users\霍冠华\AppData\Local\Programs\Python\Python313\Lib\site-packages')

try:
    import docx

    file_path = r'C:\Users\霍冠华\Desktop\简历\芦若禾简历.docx'

    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        sys.exit(1)

    # 打开文档
    doc = docx.Document(file_path)

    print("=== 芦若禾简历内容 ===\n")

    # 读取所有段落
    content = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            content.append(f"{i+1:3d}: {text}")

    # 打印内容
    for line in content:
        print(line)

    # 读取表格
    if doc.tables:
        print("\n=== 表格内容 ===")
        for t_idx, table in enumerate(doc.tables):
            print(f"\n表格 {t_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    print(f"  行 {row_idx + 1}: {' | '.join(row_data)}")

    print("\n=== 读取完成 ===")

except ImportError:
    print("错误: 无法导入python-docx库")
    print("请运行: pip install python-docx")
except Exception as e:
    print(f"读取文件时出错: {e}")
    import traceback
    traceback.print_exc()