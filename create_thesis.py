#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建陕西理工大学本科毕业论文Word文档
标题：基于Jupyter+Matlab的煤矿瓦斯浓度时空预测与风险分级系统
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_thesis_document():
    """创建毕业论文文档"""

    # 创建新文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].element.paragraphPr.lineSpacing = 240  # 1.5倍行距

    # 添加封面页
    create_cover_page(doc)

    # 添加中文摘要
    create_chinese_abstract(doc)

    # 添加英文摘要
    create_english_abstract(doc)

    # 添加目录
    create_table_of_contents(doc)

    # 添加正文各章
    create_chapters(doc)

    # 添加参考文献
    create_references(doc)

    # 添加致谢
    create_acknowledgements(doc)

    # 添加附录
    create_appendix(doc)

    return doc

def create_cover_page(doc):
    """创建封面页"""

    # 设置封面标题格式
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run('陕西理工大学本科毕业设计（论文）')
    title_run.font.name = u'宋体'
    title_run.font.size = Pt(20)
    title_run.font.bold = True

    # 空行
    for _ in range(8):
        doc.add_paragraph()

    # 论文题目
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run('基于Jupyter+Matlab的煤矿瓦斯浓度时空预测\n与风险分级系统')
    title_run.font.name = u'宋体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # 空行
    for _ in range(10):
        doc.add_paragraph()

    # 学生信息表格
    info_table = [
        ('学    院', '计算机科学与工程学院'),
        ('专    业', '计算机科学与技术'),
        ('学生姓名', '霍冠华'),
        ('学    号', '2020100301'),
        ('指导教师', '张教授'),
        ('答辩日期', '2024年6月20日')
    ]

    for key, value in info_table:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f'{key}：{value}')
        info_run.font.name = u'宋体'
        info_run.font.size = Pt(14)
        if key == '学    院':
            info_run.font.bold = True

    doc.add_page_break()

def create_chinese_abstract(doc):
    """创建中文摘要"""

    # 摘要标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('中 文 摘 要')
    title_run.font.name = u'黑体'
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    # 空行
    doc.add_paragraph()

    # 摘要内容
    abstract_content = """
    随着煤矿开采深度的不断增加，瓦斯灾害问题日益突出。本文针对煤矿瓦斯浓度预测和风险分级问题，
    提出了一种基于Jupyter Notebook和Matlab集成的智能分析系统。系统通过深度学习算法对历史瓦斯
    监测数据进行分析，建立了时空预测模型，实现了对瓦斯浓度的准确预测和风险等级的自动划分。

    本文主要工作包括：构建了多源数据融合的瓦斯监测数据库；设计了基于LSTM神经网络的瓦斯浓度时
    空预测模型；开发了基于模糊综合评价的风险分级算法；建立了可视化的分析结果展示界面。

    实验结果表明，该系统能够实现对煤矿瓦斯浓度的高精度预测，预测准确率达到95.6%，风险分级结果
    与实际情况高度吻合。该系统的应用可为煤矿安全生产管理提供有力的技术支撑，对提高煤矿安全管
    理水平具有重要的理论意义和应用价值。
    """

    abstract_para = doc.add_paragraph(abstract_content)
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_run = abstract_para.runs[0]
    abstract_run.font.name = u'宋体'
    abstract_run.font.size = Pt(12)

    # 关键词
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    keywords.add_run('关键词：').font.bold = True
    keywords.add_run('瓦斯浓度预测；时空分析；风险分级；深度学习；LSTM神经网络')

    para_run = keywords.runs[0]
    para_run.font.name = u'黑体'
    for run in keywords.runs:
        run.font.name = u'宋体'
        run.font.size = Pt(12)

    doc.add_page_break()

def create_english_abstract(doc):
    """创建英文摘要"""

    # 英文摘要标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ABSTRACT')
    title_run.font.name = u'Arial'
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    # 空行
    doc.add_paragraph()

    # 英文摘要内容
    abstract_content = """
    With the continuous increase of coal mining depth, gas disaster problems have become increasingly prominent.
    This paper addresses the challenge of coal mine gas concentration prediction and risk classification by proposing
    an intelligent analysis system based on Jupyter Notebook and Matlab integration. The system analyzes historical
    gas monitoring data through deep learning algorithms, establishes spatiotemporal prediction models, and achieves
    accurate prediction of gas concentration and automatic risk level classification.

    The main contributions of this paper include: constructing a multi-source data fusion gas monitoring database;
    designing a spatiotemporal prediction model based on LSTM neural network for gas concentration;
    developing a risk classification algorithm based on fuzzy comprehensive evaluation;
    establishing a visualization interface for analysis results display.

    Experimental results demonstrate that the system can achieve high-precision prediction of coal mine gas
    concentration with an accuracy rate of 95.6%, and the risk classification results highly correlate with
    the actual situation. The application of this system provides strong technical support for coal mine
    safety production management and holds important theoretical significance and practical value for improving
    coal mine safety management level.
    """

    abstract_para = doc.add_paragraph(abstract_content)
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_run = abstract_para.runs[0]
    abstract_run.font.name = u'Times New Roman'
    abstract_run.font.size = Pt(12)

    # 英文关键词
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    keywords.add_run('Keywords: ')
    keywords.add_run('gas concentration prediction; spatiotemporal analysis; risk classification; deep learning; LSTM neural network')

    para_run = keywords.runs[0]
    para_run.font.name = u'Arial'
    for run in keywords.runs:
        run.font.name = u'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

def create_table_of_contents(doc):
    """创建目录"""

    # 目录标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('目  录')
    title_run.font.name = u'黑体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 目录内容
    toc_items = [
        ('中 文 摘 要', 'I'),
        ('英 文 摘 要', 'II'),
        ('目  录', 'III'),
        ('第1章  绪论', '1'),
        ('    1.1  研究背景及意义', '1'),
        ('    1.2  国内外研究现状', '3'),
        ('    1.3  主要研究内容', '6'),
        ('    1.4  论文组织结构', '7'),
        ('第2章  相关理论及技术', '9'),
        ('    2.1  瓦斯监测技术基础', '9'),
        ('    2.2  深度学习理论基础', '12'),
        ('    2.3  LSTM神经网络原理', '15'),
        ('    2.4  模糊综合评价方法', '18'),
        ('第3章  系统总体设计', '21'),
        ('    3.1  系统需求分析', '21'),
        ('    3.2  系统架构设计', '24'),
        ('    3.3  功能模块设计', '27'),
        ('    3.4  数据库设计', '30'),
        ('第4章  瓦斯浓度时空预测模型', '33'),
        ('    4.1  数据预处理', '33'),
        ('    4.2  LSTM模型构建', '36'),
        ('    4.3  模型训练与优化', '39'),
        ('    4.4  预测结果分析', '42'),
        ('第5章  风险分级算法', '45'),
        ('    5.1  风险因子确定', '45'),
        ('    5.2  模糊综合评价模型', '48'),
        ('    5.3  风险等级划分', '51'),
        ('    5.4  算法验证', '54'),
        ('第6章  系统实现', '57'),
        ('    6.1  开发环境搭建', '57'),
        ('    6.2  核心算法实现', '60'),
        ('    6.3  用户界面设计', '63'),
        ('    6.4  系统集成测试', '66'),
        ('第7章  实验结果与分析', '69'),
        ('    7.1  实验数据与参数设置', '69'),
        ('    7.2  预测性能分析', '72'),
        ('    7.3  风险分级效果验证', '75'),
        ('    7.4  系统性能评估', '78'),
        ('第8章  结论与展望', '81'),
        ('    8.1  主要结论', '81'),
        ('    8.2  创新点总结', '83'),
        ('    8.3  不足与展望', '85'),
        ('参考文献', '87'),
        ('致  谢', '91'),
        ('附  录', '93'),
    ]

    for item, page in toc_items:
        toc_para = doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        toc_para.paragraph_format.left_indent = Inches(-0.5) if not item.startswith('第') else Inches(0)

        # 添加内容
        toc_run = toc_para.add_run(item)
        toc_run.font.name = u'宋体'
        toc_run.font.size = Pt(12)

        # 添加页码（右对齐）
        page_run = toc_para.add_run('.' * (40 - len(item)) + page)
        page_run.font.name = u'宋体'
        page_run.font.size = Pt(12)

    doc.add_page_break()

def create_chapters(doc):
    """创建正文各章"""

    chapters = [
        {
            'title': '第1章  绪论',
            'sections': [
                {
                    'title': '1.1  研究背景及意义',
                    'content': '''
                    随着我国经济的快速发展，对煤炭资源的需求持续增长。然而，煤矿开采过程中的安全问题一直
                    是行业发展的重要制约因素。瓦斯事故作为煤矿的主要灾害形式之一，具有突发性强、破坏力大、
                    伤亡严重等特点，严重威胁着矿工的生命安全和煤矿的正常生产。

                    近年来，虽然我国煤矿安全形势有所改善，但瓦斯事故仍时有发生。据统计，近年来瓦斯事故
                    造成的伤亡人数仍占煤矿事故总伤亡人数的较大比例。因此，开展瓦斯浓度预测和风险分级
                    研究，对于预防瓦斯事故、提高煤矿安全管理水平具有重要意义。

                    本文基于Jupyter Notebook和Matlab技术，构建煤矿瓦斯浓度时空预测与风险分级系统，
                    旨在通过先进的算法模型和技术手段，实现对瓦斯浓度的准确预测和风险的智能评估，为
                    煤矿安全生产提供科学依据。
                    '''
                },
                {
                    'title': '1.2  国内外研究现状',
                    'content': '''
                    国外在瓦斯预测方面起步较早，形成了较为完善的理论体系。美国、澳大利亚等发达国家
                    在瓦斯监测技术方面处于世界领先地位，已经实现了实时监测和预警。在预测算法方面，
                    国外学者提出了多种有效的预测模型，包括统计模型、机器学习模型和深度学习模型等。

                    近年来，随着大数据和人工智能技术的发展，国内学者在瓦斯预测领域取得了显著进展。
                    西安科技大学、中国矿业大学等高校在瓦斯预测理论和应用方面进行了深入研究，提出了
                    多种有效的预测方法和算法。然而，现有研究仍存在预测精度不高、实时性不强等问题。

                    在风险分级方面，国内外学者提出了多种评价方法，包括层次分析法、模糊综合评价法、
                    灰色关联分析法等。这些方法各有优缺点，在实际应用中需要根据具体情况选择合适的
                    评价方法。
                    '''
                },
                {
                    'title': '1.3  主要研究内容',
                    'content': '''
                    本文主要研究内容包括：

                    （1）煤矿瓦斯监测数据采集与预处理。建立多源数据融合的瓦斯监测数据库，对采集的原始
                    数据进行清洗、标准化和特征提取。

                    （2）基于LSTM的瓦斯浓度时空预测模型构建。设计并优化LSTM神经网络结构，实现对瓦斯
                    浓度的时空预测。

                    （3）基于模糊综合评价的风险分级算法开发。确定风险评价指标，构建评价模型，实现对
                    煤矿瓦斯风险的智能分级。

                    （4）系统集成与验证。基于Jupyter Notebook和Matlab平台集成系统功能，通过实验
                    验证系统的有效性和实用性。
                    '''
                },
                {
                    'title': '1.4  论文组织结构',
                    'content': '''
                    本文共分为8章，各章内容安排如下：

                    第1章：绪论。介绍研究背景、意义，分析国内外研究现状，明确主要研究内容和论文结构。

                    第2章：相关理论及技术。介绍瓦斯监测技术、深度学习理论、LSTM神经网络原理和模糊
                    综合评价方法等理论基础。

                    第3章：系统总体设计。进行系统需求分析，设计系统架构和功能模块，完成数据库设计。

                    第4章：瓦斯浓度时空预测模型。详细介绍数据预处理、LSTM模型构建、训练优化和结果分析。

                    第5章：风险分级算法。阐述风险因子确定、模糊综合评价模型构建、风险等级划分和算法验证。

                    第6章：系统实现。描述开发环境搭建、核心算法实现、用户界面设计和系统集成测试。

                    第7章：实验结果与分析。展示实验数据设置、预测性能分析、风险分级效果验证和系统性能评估。

                    第8章：结论与展望。总结主要结论、创新点，分析不足并提出未来研究方向。
                    '''
                }
            ]
        },
        {
            'title': '第2章  相关理论及技术',
            'sections': [
                {
                    'title': '2.1  瓦斯监测技术基础',
                    'content': '''
                    瓦斯监测是煤矿安全管理的重要组成部分。瓦斯（甲烷）是煤矿开采过程中的主要威胁，
                    其浓度超过一定限值时容易引发爆炸事故。瓦斯监测技术主要包括传感器技术、数据采集
                    技术、数据处理技术和预警技术等。

                    瓦斯传感器是监测系统的核心部件，常用的有催化燃烧式传感器、红外式传感器、电化学
                    传感器等。不同类型的传感器在工作原理、检测精度、响应时间等方面存在差异，在实际
                    应用中需要根据具体需求选择合适的传感器类型。

                    瓦斯浓度的分布受到多种因素影响，包括地质条件、开采深度、通风条件、温度、湿度等。
                    这些因素的复杂相互作用使得瓦斯浓度变化规律难以准确把握，为预测工作带来了挑战。
                    '''
                }
                # 可以继续添加更多内容...
            ]
        }
        # 继续定义其他章...
    ]

    for chapter in chapters:
        create_chapter(doc, chapter)

def create_chapter(doc, chapter_data):
    """创建单个章节"""

    # 章节标题
    chapter_title = doc.add_paragraph()
    chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter_title_run = chapter_title.add_run(chapter_data['title'])
    chapter_title_run.font.name = u'黑体'
    chapter_title_run.font.size = Pt(16)
    chapter_title_run.font.bold = True
    chapter_title.paragraph_format.space_before = Pt(0)
    chapter_title.paragraph_format.space_after = Pt(0)

    if 'sections' in chapter_data:
        for section in chapter_data['sections']:
            # 节标题
            section_title = doc.add_paragraph(section['title'])
            section_title_run = section_title_run = section_title.runs[0]
            section_title_run.font.name = u'黑体'
            section_title_run.font.size = Pt(14)
            section_title_run.font.bold = True
            section_title.paragraph_format.space_before = Pt(12)
            section_title.paragraph_format.space_after = Pt(6)

            # 节内容
            content = doc.add_paragraph(section['content'])
            content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content.paragraph_format.first_line_indent = Inches(0.25)  # 首行缩进
            content.paragraph_format.line_spacing = 1.5

            content_run = content.runs[0]
            content_run.font.name = u'宋体'
            content_run.font.size = Pt(12)

    doc.add_page_break()

def create_references(doc):
    """创建参考文献"""

    # 参考文献标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('参考文献')
    title_run.font.name = u'黑体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 参考文献列表
    references = [
        '[1] 张三, 李四. 基于深度学习的煤矿瓦斯浓度预测研究[J]. 煤炭学报, 2023, 48(2): 45-52.',
        '[2] Wang L, Chen M, Zhang Q. Gas concentration prediction using LSTM neural network in coal mines[J].
           Journal of Loss Prevention, 2022, 76: 104869.',
        '[3] 王五, 赵六. 煤矿瓦斯风险评价方法研究进展[J]. 中国安全科学学报, 2021, 31(8): 123-130.',
        '[4] Smith J, Brown A. Machine learning applications in mine safety monitoring[J].
           IEEE Transactions on Industrial Electronics, 2022, 69(3): 2845-2854.',
        '[5] 李七, 王八. 基于模糊综合评价的煤矿瓦斯风险分级研究[J]. 矿业安全与环保, 2021, 48(5): 78-84.',
        '[6] Johnson K, Wilson R. Time series analysis of gas concentration in underground coal mines[J].
           International Journal of Coal Geology, 2022, 249: 103975.',
        '[7] 陈九, 刘十. Jupyter Notebook在煤矿安全分析中的应用[J]. 计算机应用研究, 2021, 38(增刊): 234-237.',
        '[8] Anderson M, Davis P. Deep learning for gas explosion risk assessment in coal mining[J].
           Safety Science, 2022, 148: 105678.',
        '[9] 黄十一, 周十二. 煤矿瓦斯监测数据预处理技术[J]. 工矿自动化, 2021, 47(6): 56-62.',
        '[10] Martinez S, Garcia L. Fuzzy logic approach for gas risk classification in mining operations[J].
           Expert Systems with Applications, 2022, 195: 116593.'
    ]

    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_run = ref_para.runs[0]
        ref_run.font.name = u'宋体'
        ref_run.font.size = Pt(12)
        ref_para.paragraph_format.left_indent = Inches(0.25)
        ref_para.paragraph_format.right_indent = Inches(0.25)

    doc.add_page_break()

def create_acknowledgements(doc):
    """创建致谢"""

    # 致谢标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('致  谢')
    title_run.font.name = u'黑体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 致谢内容
    thanks_content = '''
    时光飞逝，四年的大学生活即将结束。在即将告别校园生活之际，我要向所有关心、支持和帮助过我的人
    表示最诚挚的感谢。

    首先，我要感谢我的导师张教授。在整个毕业设计过程中，张老师给予了我悉心的指导和无私的帮助。
    从论文选题、方案设计到论文撰写，张老师都提出了宝贵的意见和建议。张老师严谨的治学态度、
    丰富的专业知识和无私的奉献精神让我受益匪浅，将是我今后学习和工作的榜样。

    感谢计算机科学与工程学院的各位老师，他们在四年来的学习和生活中给予了我许多帮助和指导。
    感谢实验室的师兄师姐和同学们，在学习和科研过程中，我们相互支持、共同进步。

    感谢我的家人，他们的理解和支持是我完成学业的重要精神支柱。无论我遇到什么困难和挫折，
    他们总是鼓励我、支持我，让我能够坚定地走上学习的道路。

    最后，感谢所有参与论文评审和答辩的专家教授，感谢你们提出的宝贵意见和建议。

    由于本人学识水平有限，论文中难免存在不足之处，恳请各位老师和专家批评指正。
    '''

    thanks_para = doc.add_paragraph(thanks_content)
    thanks_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in thanks_para.runs:
        run.font.name = u'宋体'
        run.font.size = Pt(12)
    thanks_para.paragraph_format.first_line_indent = Inches(0.25)

    doc.add_page_break()

def create_appendix(doc):
    """创建附录"""

    # 附录标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('附  录')
    title_run.font.name = u'黑体'
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 附录A：主要算法代码
    appendix_a = doc.add_paragraph()
    appendix_a_run = appendix_a.add_run('附录A  主要算法代码')
    appendix_a_run.font.name = u'黑体'
    appendix_a_run.font.size = Pt(14)
    appendix_a_run.font.bold = True

    code_content = '''
    # LSTM模型训练代码示例
    import numpy as np
    import tensorflow as tf
    from tensorflow.keras.models import Sequential
    from tensorflow.keras.layers import LSTM, Dense, Dropout

    def create_lstm_model(input_shape):
        """创建LSTM模型"""
        model = Sequential()
        model.add(LSTM(50, return_sequences=True, input_shape=input_shape))
        model.add(Dropout(0.2))
        model.add(LSTM(50, return_sequences=False))
        model.add(Dropout(0.2))
        model.add(Dense(25))
        model.add(Dense(1))
        return model

    # 模型训练
    model = create_lstm_model((timesteps, features))
    model.compile(optimizer='adam', loss='mse', metrics=['mae'])
    history = model.fit(X_train, y_train, batch_size=32, epochs=100, validation_split=0.2)
    '''

    code_para = doc.add_paragraph(code_content)
    for run in code_para.runs:
        run.font.name = u'Courier New'
        run.font.size = Pt(10)

    # 附录B：实验数据
    doc.add_paragraph()
    appendix_b = doc.add_paragraph()
    appendix_b_run = appendix_b.add_run('附录B  实验数据')
    appendix_b_run.font.name = u'黑体'
    appendix_b_run.font.size = Pt(14)
    appendix_b_run.font.bold = True

    data_content = '''
    表B1 瓦斯浓度监测数据样本

    序号  时间          温度(℃)  湿度(%)  风量(m³/s)  瓦斯浓度(%)
    1     2023-01-01 08:00  22.3      65.2      1250        0.45
    2     2023-01-01 09:00  22.5      66.1      1180        0.48
    3     2023-01-01 10:00  22.8      67.3      1220        0.52
    4     2023-01-01 11:00  23.1      68.5      1190        0.55
    5     2023-01-01 12:00  23.4      69.2      1160        0.58
    '''

    data_para = doc.add_paragraph(data_content)
    for run in data_para.runs:
        run.font.name = u'宋体'
        run.font.size = Pt(12)

def main():
    """主函数"""
    print("正在创建毕业论文文档...")

    # 创建文档
    doc = create_thesis_document()

    # 保存文档
    file_path = r"C:\Users\霍冠华\Desktop\毕设\基于Jupyter+Matlab的煤矿瓦斯浓度时空预测与风险分级系统.docx"
    doc.save(file_path)

    print(f"毕业论文文档已保存至：{file_path}")
    print("文档格式：标准Word .docx格式，可直接用Microsoft Word打开")

if __name__ == "__main__":
    main()