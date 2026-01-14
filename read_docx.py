#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
import sys

def read_docx_file(file_path):
    try:
        doc = docx.Document(file_path)
        print("=== 芦若禾简历内容 ===")
        print()

        # 读取所有段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"{i+1:3d}: {text}")

        # 读取表格内容
        if doc.tables:
            print("\n=== 表格内容 ===")
            for t_idx, table in enumerate(doc.tables):
                print(f"\n表格 {t_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        print(f"  行 {row_idx + 1}: {' | '.join(row_text)}")

    except Exception as e:
        print(f"读取文件时出错: {e}")
        return False
    return True

if __name__ == "__main__":
    file_path = r"C:\Users\霍冠华\Desktop\简历\芦若禾简历.docx"
    read_docx_file(file_path)
